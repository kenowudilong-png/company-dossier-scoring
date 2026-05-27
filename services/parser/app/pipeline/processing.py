from __future__ import annotations

import hashlib
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import BinaryIO

try:
    import pdfplumber
except Exception:  # pragma: no cover
    pdfplumber = None
try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None
try:
    import openpyxl
except Exception:  # pragma: no cover
    openpyxl = None

@dataclass(frozen=True)
class ParsedPage:
    page_number: int
    text: str

@dataclass(frozen=True)
class Chunk:
    chunk_id: str
    page_number: int
    content: str
    snippet: str


def stable_vector(text: str, dim: int = 1536) -> list[float]:
    digest = hashlib.sha256(text.encode()).digest()
    return [round((digest[i % len(digest)] / 255.0 - 0.5) / dim, 8) for i in range(dim)]


def vector_literal(values: list[float]) -> str:
    return "[" + ",".join(str(value) for value in values) + "]"


def parse_txt(data: bytes) -> list[ParsedPage]:
    return [ParsedPage(1, data.decode("utf-8", errors="ignore"))]


def parse_pdf(data: bytes) -> list[ParsedPage]:
    if pdfplumber is None:
        return [ParsedPage(1, "")]
    pages = []
    with pdfplumber.open(BytesIO(data)) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            tables = page.extract_tables() or []
            table_text = "\n".join(" | ".join(str(cell or "") for cell in row) for table in tables for row in table)
            pages.append(ParsedPage(index, "\n".join(part for part in [text, table_text] if part).strip()))
    return pages or [ParsedPage(1, "")]


def parse_docx(data: bytes) -> list[ParsedPage]:
    if Document is None:
        return [ParsedPage(1, "")]
    doc = Document(BytesIO(data))
    parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return [ParsedPage(1, "\n".join(parts))]


def parse_xlsx(data: bytes) -> list[ParsedPage]:
    if openpyxl is None:
        return [ParsedPage(1, "")]
    wb = openpyxl.load_workbook(BytesIO(data), data_only=True)
    pages = []
    for index, sheet in enumerate(wb.worksheets, start=1):
        rows = []
        for row in sheet.iter_rows(values_only=True):
            values = [str(value) for value in row if value is not None]
            if values:
                rows.append(" | ".join(values))
        pages.append(ParsedPage(index, f"Sheet: {sheet.title}\n" + "\n".join(rows)))
    return pages or [ParsedPage(1, "")]


def parse_file(filename: str, mime_type: str, data: bytes) -> list[ParsedPage]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf" or "pdf" in mime_type:
        return parse_pdf(data)
    if suffix in {".docx", ".doc"} or "word" in mime_type:
        return parse_docx(data)
    if suffix in {".xlsx", ".xls"} or "spreadsheet" in mime_type or "excel" in mime_type:
        return parse_xlsx(data)
    return parse_txt(data)


def chunk_pages(file_id: str, pages: list[ParsedPage], max_chars: int = 900) -> list[Chunk]:
    chunks = []
    for page in pages:
        text = " ".join(page.text.split())
        if not text:
            continue
        start = 0
        while start < len(text):
            content = text[start:start + max_chars].strip()
            digest = hashlib.sha1(f"{file_id}:{page.page_number}:{content}".encode()).hexdigest()[:16]
            chunks.append(Chunk(f"chk_{digest}", page.page_number, content, content[:220]))
            start += max_chars
    return chunks
